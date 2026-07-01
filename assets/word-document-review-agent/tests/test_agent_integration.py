"""Integration tests for the WordReviewAgent."""
import sys
import os
import json
import tempfile
from pathlib import Path
from unittest.mock import AsyncMock, MagicMock, patch

APP_DIR = Path(__file__).parent.parent / "app"
if str(APP_DIR) not in sys.path:
    sys.path.insert(0, str(APP_DIR))

import pytest
from docx import Document


def _make_fixture_docx(text: str = "This document has a misteake and bad grammer.") -> str:
    """Create a fixture .docx file for testing."""
    fixtures_dir = Path(__file__).parent / "fixtures"
    fixtures_dir.mkdir(exist_ok=True)
    path = str(fixtures_dir / "test_doc.docx")
    doc = Document()
    doc.add_paragraph(text)
    doc.add_paragraph("Another paragraph with no issues.")
    doc.save(path)
    return path


MOCK_ISSUES = [
    {
        "paragraph_index": 0,
        "original_text": "misteake",
        "issue_type": "spelling",
        "explanation": "'misteake' should be 'mistake'",
        "suggestion": "mistake",
    },
    {
        "paragraph_index": 0,
        "original_text": "bad grammer",
        "issue_type": "grammar",
        "explanation": "'grammer' should be 'grammar'",
        "suggestion": "grammar",
    },
]

MOCK_LLM_RESPONSE = json.dumps(MOCK_ISSUES)


@pytest.mark.asyncio
async def test_agent_invoke_returns_completed():
    """Full end-to-end: agent receives a file path, reviews it, returns completed."""
    from agent import WordReviewAgent

    file_path = _make_fixture_docx()

    # Mock the LLM response
    mock_graph = AsyncMock()
    mock_graph.ainvoke.return_value = {
        "messages": [MagicMock(content=f"Here are the issues:\n{MOCK_LLM_RESPONSE}")]
    }

    with patch("langchain.agents.create_agent", return_value=mock_graph):
        agent = WordReviewAgent()
        result = await agent.invoke(file_path, context_id="test-context-1", tools=[])

    assert result.status == "completed"
    assert "Review complete" in result.message or "issue" in result.message.lower()


@pytest.mark.asyncio
async def test_agent_handles_missing_file():
    """Agent should return error message for missing file."""
    from agent import WordReviewAgent

    agent = WordReviewAgent()
    result = await agent.invoke("/nonexistent/path/missing.docx", context_id="test-context-2", tools=[])

    assert result.status == "completed"
    assert "REVIEW ERROR" in result.message or "not found" in result.message.lower()


@pytest.mark.asyncio
async def test_agent_handles_non_docx():
    """Agent should return error message for non-.docx files."""
    from agent import WordReviewAgent

    with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as tmp:
        tmp.write(b"not a docx")
        path = tmp.name

    try:
        agent = WordReviewAgent()
        result = await agent.invoke(path, context_id="test-context-3", tools=[])
        assert result.status == "completed"
        assert "REVIEW ERROR" in result.message or "not a .docx" in result.message.lower()
    finally:
        os.unlink(path)


@pytest.mark.asyncio
async def test_agent_stream_yields_progress_then_result():
    """Agent stream should yield a 'processing' update before the final result."""
    from agent import WordReviewAgent

    file_path = _make_fixture_docx()
    mock_graph = AsyncMock()
    mock_graph.ainvoke.return_value = {
        "messages": [MagicMock(content=f"[]\n{MOCK_LLM_RESPONSE}")]
    }

    with patch("langchain.agents.create_agent", return_value=mock_graph):
        agent = WordReviewAgent()
        chunks = []
        async for chunk in agent.stream(file_path, "test-context-4", tools=[]):
            chunks.append(chunk)

    # First chunk should be a progress update
    assert len(chunks) >= 2
    assert chunks[0]["is_task_complete"] is False
    # Last chunk should be complete
    assert chunks[-1]["is_task_complete"] is True


@pytest.mark.asyncio
async def test_agent_empty_json_response_graceful():
    """Agent handles empty issue list from LLM gracefully."""
    from agent import WordReviewAgent

    file_path = _make_fixture_docx()
    mock_graph = AsyncMock()
    mock_graph.ainvoke.return_value = {
        "messages": [MagicMock(content="No issues found. []")]
    }

    with patch("langchain.agents.create_agent", return_value=mock_graph):
        agent = WordReviewAgent()
        result = await agent.invoke(file_path, context_id="test-context-5", tools=[])

    assert result.status == "completed"
    assert "0 issue" in result.message or "issue" in result.message.lower()
