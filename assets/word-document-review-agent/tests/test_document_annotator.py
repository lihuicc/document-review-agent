"""Unit tests for document_annotator tool."""
import sys
import os
import tempfile
from pathlib import Path

APP_DIR = Path(__file__).parent.parent / "app"
if str(APP_DIR) not in sys.path:
    sys.path.insert(0, str(APP_DIR))

import pytest
from docx import Document

from tools.document_annotator import annotate_document


def _make_temp_docx(paragraphs: list[str]) -> str:
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(tmp.name)
    tmp.close()
    return tmp.name


def test_annotate_document_creates_output_file():
    src = _make_temp_docx(["This is a test sentence with a misteake."])
    out = src.replace(".docx", "-reviewed.docx")
    issues = [{
        "paragraph_index": 0,
        "original_text": "misteake",
        "issue_type": "spelling",
        "explanation": "'misteake' is misspelled",
        "suggestion": "mistake",
    }]
    try:
        result = annotate_document(src, issues, out)
        assert Path(result).exists()
    finally:
        os.unlink(src)
        if Path(out).exists():
            os.unlink(out)


def test_annotate_document_returns_output_path():
    src = _make_temp_docx(["Hello world."])
    out = src.replace(".docx", "-reviewed.docx")
    issues = [{
        "paragraph_index": 0,
        "issue_type": "clarity",
        "explanation": "Vague greeting",
        "suggestion": "Be more specific",
    }]
    try:
        result = annotate_document(src, issues, out)
        assert result == out
    finally:
        os.unlink(src)
        if Path(out).exists():
            os.unlink(out)


def test_annotate_document_no_issues_saves_copy():
    src = _make_temp_docx(["Perfect sentence."])
    out = src.replace(".docx", "-reviewed.docx")
    try:
        result = annotate_document(src, [], out)
        assert Path(result).exists()
        doc = Document(result)
        texts = [p.text for p in doc.paragraphs]
        assert "Perfect sentence." in texts
    finally:
        os.unlink(src)
        if Path(out).exists():
            os.unlink(out)


def test_annotate_document_out_of_range_paragraph():
    """Should not crash when paragraph_index exceeds document length."""
    src = _make_temp_docx(["Only one paragraph."])
    out = src.replace(".docx", "-reviewed.docx")
    issues = [{
        "paragraph_index": 999,
        "issue_type": "grammar",
        "explanation": "Issue at non-existent paragraph",
        "suggestion": "Fix it",
    }]
    try:
        result = annotate_document(src, issues, out)
        assert Path(result).exists()
    finally:
        os.unlink(src)
        if Path(out).exists():
            os.unlink(out)


def test_annotate_document_source_not_found():
    with pytest.raises(ValueError, match="not found"):
        annotate_document("/nonexistent/file.docx", [], "/tmp/out.docx")


def test_annotate_document_preserves_original_content():
    src = _make_temp_docx(["Original text here.", "Second line."])
    out = src.replace(".docx", "-reviewed.docx")
    issues = [{
        "paragraph_index": 0,
        "issue_type": "spelling",
        "explanation": "Test issue",
        "suggestion": "Fix",
    }]
    try:
        annotate_document(src, issues, out)
        doc = Document(out)
        texts = [p.text for p in doc.paragraphs]
        assert "Original text here." in texts
        assert "Second line." in texts
    finally:
        os.unlink(src)
        if Path(out).exists():
            os.unlink(out)
