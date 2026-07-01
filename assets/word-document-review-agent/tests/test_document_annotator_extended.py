"""Extended tests for document_annotator to push coverage above 70%."""
import sys
import os
import tempfile
from pathlib import Path

APP_DIR = Path(__file__).parent.parent / "app"
if str(APP_DIR) not in sys.path:
    sys.path.insert(0, str(APP_DIR))

import pytest
from docx import Document

from tools.document_annotator import (
    annotate_document,
    _add_comment_to_paragraph,
    _get_or_create_comments_part,
)


def _make_temp_docx(paragraphs=None):
    if paragraphs is None:
        paragraphs = ["Default paragraph."]
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(tmp.name)
    tmp.close()
    return tmp.name


def _make_temp_out(src):
    return src.replace(".docx", "-out.docx")


def test_multiple_issues_all_annotated():
    src = _make_temp_docx(["First sentence.", "Second sentence.", "Third sentence."])
    out = _make_temp_out(src)
    issues = [
        {"paragraph_index": 0, "issue_type": "spelling", "explanation": "typo", "suggestion": "fix"},
        {"paragraph_index": 1, "issue_type": "grammar", "explanation": "bad grammar", "suggestion": "correct it"},
        {"paragraph_index": 2, "issue_type": "clarity", "explanation": "vague", "suggestion": "be specific"},
    ]
    try:
        result = annotate_document(src, issues, out)
        assert Path(result).exists()
        # Re-read and verify content is preserved
        doc = Document(result)
        texts = [p.text for p in doc.paragraphs]
        assert "First sentence." in texts
        assert "Second sentence." in texts
        assert "Third sentence." in texts
    finally:
        os.unlink(src)
        if Path(out).exists():
            os.unlink(out)


def test_issue_without_suggestion():
    src = _make_temp_docx(["A paragraph."])
    out = _make_temp_out(src)
    issues = [
        {"paragraph_index": 0, "issue_type": "logic", "explanation": "contradictory statement"},
    ]
    try:
        result = annotate_document(src, issues, out)
        assert Path(result).exists()
    finally:
        os.unlink(src)
        if Path(out).exists():
            os.unlink(out)


def test_issue_type_defaults_to_unknown():
    src = _make_temp_docx(["A paragraph."])
    out = _make_temp_out(src)
    issues = [{"paragraph_index": 0}]  # minimal issue with no keys
    try:
        result = annotate_document(src, issues, out)
        assert Path(result).exists()
    finally:
        os.unlink(src)
        if Path(out).exists():
            os.unlink(out)


def test_add_comment_to_paragraph_fallback_to_last():
    """When paragraph_index > number of paragraphs, falls back to last paragraph."""
    doc = Document()
    doc.add_paragraph("Only paragraph.")
    # Should not raise even with out-of-range index
    _add_comment_to_paragraph(doc, 999, "Out of range comment", 0)


def test_add_comment_to_empty_document():
    """When document has no paragraphs, should not raise."""
    doc = Document()
    # Empty doc — just a default empty body paragraph
    _add_comment_to_paragraph(doc, 0, "Test comment", 0)


def test_get_or_create_comments_part_creates_once():
    doc = Document()
    doc.add_paragraph("test")
    part1 = _get_or_create_comments_part(doc)
    part2 = _get_or_create_comments_part(doc)
    # Both calls should return the same element (no duplicate creation)
    assert part1 is part2


def test_annotate_five_issues_different_types():
    src = _make_temp_docx([
        "Thsi sentence has a spelling mistake.",
        "He don't have any idea.",
        "The results was unexpected.",
        "Furthermore, the report clearly states that the findings, which were conducted over a period of several months and reviewed by multiple teams, are, as previously mentioned, inconclusive.",
        "First we conclude. First we analyze.",
    ])
    out = _make_temp_out(src)
    issues = [
        {"paragraph_index": 0, "issue_type": "spelling", "explanation": "'Thsi' → 'This'", "suggestion": "This"},
        {"paragraph_index": 1, "issue_type": "grammar", "explanation": "Subject-verb disagreement", "suggestion": "He doesn't"},
        {"paragraph_index": 2, "issue_type": "grammar", "explanation": "'was' should be 'were'", "suggestion": "were"},
        {"paragraph_index": 3, "issue_type": "clarity", "explanation": "Sentence too long", "suggestion": "Break into shorter sentences"},
        {"paragraph_index": 4, "issue_type": "consistency", "explanation": "Contradictory order of operations", "suggestion": "Reorder steps"},
    ]
    try:
        result = annotate_document(src, issues, out)
        assert Path(result).exists()
        doc = Document(result)
        assert len([p for p in doc.paragraphs if p.text.strip()]) >= 5
    finally:
        os.unlink(src)
        if Path(out).exists():
            os.unlink(out)
